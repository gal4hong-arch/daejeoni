"""평문·간단 마크다운(# 제목)을 .docx 바이트로 변환."""

from __future__ import annotations

from io import BytesIO


def text_to_docx_bytes(text: str, title: str | None = None) -> bytes:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx 패키지가 필요합니다. 저장소 루트에서 poetry install 후 서버를 다시 실행하세요."
        ) from e

    doc = Document()
    tit = (title or "").strip()
    if tit:
        doc.add_heading(tit[:200], 0)

    lines = (text or "").replace("\r\n", "\n").split("\n")
    for line in lines:
        s = line.rstrip()
        if not s:
            doc.add_paragraph("")
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=2)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=1)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=0)
        else:
            doc.add_paragraph(s)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
