"""보고서 양식 파일(HWP/HWPX/DOC/DOCX) → 평문 추출(작성 시 구조 힌트용)."""

from __future__ import annotations

import io
from pathlib import Path

MAX_TEMPLATE_BYTES = 8 * 1024 * 1024
ALLOWED_EXT = frozenset({".docx", ".doc", ".hwp", ".hwpx"})


def extract_template_plaintext(filename: str, data: bytes) -> str:
    if not data:
        raise ValueError("빈 파일입니다.")
    if len(data) > MAX_TEMPLATE_BYTES:
        raise ValueError(f"양식 파일이 너무 큽니다(최대 약 {MAX_TEMPLATE_BYTES // (1024 * 1024)}MB).")
    name = (filename or "").strip() or "template"
    ext = Path(name).suffix.lower()
    if ext not in ALLOWED_EXT:
        raise ValueError(f"지원 양식: {', '.join(sorted(ALLOWED_EXT))}")

    if ext == ".docx":
        return _from_docx(data)
    return _from_unstructured(name, data)


def _from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("DOCX 처리에 python-docx가 필요합니다.") from e
    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [((c.text or "").strip()) for c in row.cells]
            row_txt = " | ".join(x for x in cells if x)
            if row_txt:
                lines.append(row_txt)
    out = "\n".join(lines).strip()
    if len(out) < 20:
        raise ValueError("DOCX에서 추출한 텍스트가 거의 없습니다. 표·머리글이 비었을 수 있습니다.")
    return out[:48000]


def _from_unstructured(filename: str, data: bytes) -> str:
    try:
        from unstructured.partition.auto import partition
    except ImportError as e:
        raise ValueError("이 형식 처리에 unstructured 패키지가 필요합니다.") from e
    bio = io.BytesIO(data)
    bio.seek(0)
    try:
        elements = partition(file=bio, metadata_filename=filename)
    except Exception as e:
        raise ValueError(
            f"양식 본문을 추출하지 못했습니다. DOCX로 저장 후 다시 시도하세요. ({e})"
        ) from e
    parts = [str(el).strip() for el in elements if str(el).strip()]
    out = "\n".join(parts).strip()
    if len(out) < 20:
        raise ValueError("양식에서 읽을 텍스트가 거의 없습니다.")
    return out[:48000]
